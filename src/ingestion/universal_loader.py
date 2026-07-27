import os
import json
import PyPDF2
from docx import Document as DocxDocument
from langchain_core.documents import Document

class UniversalLoader:
    """
    The MjuziAI Universal Loader.
    Loads and normalizes text from PDF, TXT, JSON, and DOCX files.
    """
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load_document(self) -> list[Document]:
        _, ext = os.path.splitext(self.file_path.lower())
        
        if ext == ".pdf":
            return self._load_pdf()
        elif ext == ".json":
            return self._load_json()
        elif ext == ".txt":
            return self._load_txt()
        elif ext == ".docx":
            return self._load_docx()
        else:
            raise ValueError(f"❌ Unsupported file format: {ext}")

    def _load_pdf(self) -> list[Document]:
        docs = []
        with open(self.file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                docs.append(Document(page_content=text, metadata={"source": self.file_path, "page": i+1}))
        return docs

    def _load_txt(self) -> list[Document]:
        with open(self.file_path, "r", encoding="utf-8") as f:
            return [Document(page_content=f.read(), metadata={"source": self.file_path})]

    def _load_docx(self) -> list[Document]:
        doc = DocxDocument(self.file_path)
        full_text = [para.text for para in doc.paragraphs]
        return [Document(page_content="\n".join(full_text), metadata={"source": self.file_path})]

    def _load_json(self) -> list[Document]:
        documents = []
        with open(self.file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
            def process_item(item):
                if isinstance(item, dict):
                    # Logic for our specialized Cultural Datasets
                    if "phrase" in item and "m.s." in item:
                        page_content = f"Tashbihi: {item['phrase']} | Maana: {item['m.s.']} | Mfano: {item['mf']}"
                    elif "proverb" in item:
                        page_content = f"Methali: {item['proverb']} | Maana: {item['m.s.']}"
                    else:
                        page_content = json.dumps(item, indent=2, ensure_ascii=False)
                    
                    return Document(page_content=page_content, metadata=item.get("metadata", {"source": self.file_path}))
                return Document(page_content=str(item), metadata={"source": self.file_path})

            if isinstance(data, list):
                for item in data:
                    documents.append(process_item(item))
            else:
                documents.append(process_item(data))
        return documents